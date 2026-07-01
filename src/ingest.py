"""
Document ingestion module.

Responsible for:
1. Reading resumes and job descriptions.
2. Extracting text from PDF, DOCX and TXT files.
3. Splitting text into chunks for RAG.
"""

from pathlib import Path
from typing import List, Dict

import docx
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

from src.config import settings


# ==========================================================
# File Readers
# ==========================================================

def _read_pdf(path: Path) -> str:
    """Extract text from a PDF file."""
    try:
        reader = PdfReader(str(path))

        text = []
        for page in reader.pages:
            text.append(page.extract_text() or "")

        return "\n".join(text)

    except Exception as e:
        raise RuntimeError(f"Failed to read PDF '{path.name}': {e}")


def _read_docx(path: Path) -> str:
    """Extract text from a DOCX file."""
    try:
        document = docx.Document(str(path))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)

    except Exception as e:
        raise RuntimeError(f"Failed to read DOCX '{path.name}': {e}")


def _read_txt(path: Path) -> str:
    """Extract text from TXT or Markdown."""
    try:
        return path.read_text(
            encoding="utf-8",
            errors="ignore"
        )

    except Exception as e:
        raise RuntimeError(f"Failed to read TXT '{path.name}': {e}")


# ==========================================================
# Public Loader
# ==========================================================

def load_text_from_file(path: str) -> str:
    """
    Load text from a supported document.

    Supported:
    - PDF
    - DOCX
    - TXT
    - MD
    """

    file_path = Path(path)

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return _read_pdf(file_path)

    if suffix == ".docx":
        return _read_docx(file_path)

    if suffix in [".txt", ".md"]:
        return _read_txt(file_path)

    raise ValueError(
        f"Unsupported file format: {suffix}"
    )


# ==========================================================
# Chunking
# ==========================================================

def chunk_text(
    text: str,
    metadata: Dict
) -> List[Document]:
    """
    Split long text into overlapping chunks.
    """

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        separators=[
            "\n\n",
            "\n",
            ". ",
            " ",
            ""
        ]
    )

    chunks = splitter.split_text(text)

    return [
        Document(
            page_content=chunk,
            metadata=metadata
        )
        for chunk in chunks
    ]


# ==========================================================
# Build Documents
# ==========================================================

def build_documents(
    resume_path: str,
    jd_text_or_path: str,
    jd_is_file: bool = False,
) -> List[Document]:
    """
    Convert resume and job description into
    LangChain Documents.
    """

    resume_text = load_text_from_file(resume_path)

    if jd_is_file:
        jd_text = load_text_from_file(jd_text_or_path)
    else:
        jd_text = jd_text_or_path

    documents = []

    documents.extend(
        chunk_text(
            resume_text,
            {
                "source": "resume",
                "filename": Path(resume_path).name,
            },
        )
    )

    documents.extend(
        chunk_text(
            jd_text,
            {
                "source": "job_description",
                "filename": (
                    Path(jd_text_or_path).name
                    if jd_is_file
                    else "manual_input"
                ),
            },
        )
    )

    return documents